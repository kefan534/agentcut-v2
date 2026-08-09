"""P0: 文档解析 + OCR 服务

能力：
- 文档类（pdf / docx / xlsx / xls / txt / md / csv）：纯文本提取
- 图片类（jpg / png / webp / gif）：腾讯云智能结构化识别 OCR
- 视频 / 音频：暂不解析（提取元数据足够）

约束（与 PRD v1.5 §3.1 对齐）：
- 单文件 ≤ 50MB
- 解析文本 ≤ 30,000 字符（超出截断并标记）
- 不可信数据使用，不进入指令上下文
- 任何异常不抛给用户，返回 text_status='failed' + text_error
"""
from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

MAX_FILE_BYTES = 50 * 1024 * 1024  # 50MB
MAX_TEXT_CHARS = 30_000  # 30k chars


# 文档扩展名 -> 解析器
def _extract_pdf(path: Path) -> Tuple[str, dict]:
    import pdfplumber  # type: ignore
    pages = []
    page_count = 0
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                pages.append(text)
    return ("\n\n".join(pages)), {"page_count": page_count}


def _extract_docx(path: Path) -> Tuple[str, dict]:
    import docx  # type: ignore
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return ("\n".join(paragraphs)), {"paragraph_count": len(doc.paragraphs)}


def _extract_xlsx(path: Path) -> Tuple[str, dict]:
    import openpyxl  # type: ignore
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            sheets.append(f"## Sheet: {ws.title}\n" + "\n".join(rows))
    return ("\n\n".join(sheets)), {"sheet_count": len(wb.worksheets)}


def _extract_text_plain(path: Path) -> Tuple[str, dict]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return text, {"line_count": text.count("\n") + 1}


def _extract_ocr_tencent(path: Path) -> Tuple[str, dict]:
    """调用腾讯云 OCR（含图片压缩）。"""
    from app.core.config import settings
    secret_id = settings.TENCENT_OCR_SECRET_ID or ""
    secret_key = settings.TENCENT_OCR_SECRET_KEY or ""
    if not secret_id or not secret_key:
        raise RuntimeError("腾讯云 OCR 凭据未配置 (TENCENT_OCR_SECRET_ID / TENCENT_OCR_SECRET_KEY)")

    # 压缩大图，避免 OCR 失败
    from io import BytesIO
    try:
        from PIL import Image  # type: ignore
        with Image.open(path) as img:
            w, h = img.size
            if max(w, h) > 4096:
                ratio = 4096 / max(w, h)
                resized = img.resize((int(w * ratio), int(h * ratio)))
            else:
                resized = img.copy()
        try:
            buf = BytesIO()
            fmt = (resized.format or "PNG").upper()
            if fmt == "JPG":
                fmt = "JPEG"
            resized.save(buf, format=fmt)
            data = buf.getvalue()
        finally:
            if resized is not img:
                resized.close()
    except Exception:
        data = path.read_bytes()

    from tencentcloud.common import credential  # type: ignore
    from tencentcloud.common.exception.tencent_cloud_sdk_exception import TencentCloudSDKException  # type: ignore
    from tencentcloud.ocr.v20181119 import ocr_client, models  # type: ignore

    cred = credential.Credential(secret_id, secret_key)
    client = ocr_client.OcrClient(cred, settings.TENCENT_OCR_REGION)
    req = models.GeneralBasicOCRRequest()
    import base64
    req.ImageBase64 = base64.b64encode(data).decode("utf-8")
    req.LanguageType = "zh"
    try:
        resp = client.GeneralBasicOCR(req)
    except TencentCloudSDKException as sdk_err:
        raise RuntimeError(f"腾讯云 OCR 调用失败: {sdk_err.get_msg()} (code={sdk_err.get_code()})") from sdk_err
    parts = [d.DetectedText for d in (resp.TextDetections or []) if d.DetectedText]
    return ("\n".join(parts)), {"service": "tencent-ocr"}


DOC_PARSERS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".xls": _extract_xlsx,
    ".txt": _extract_text_plain,
    ".md": _extract_text_plain,
    ".csv": _extract_text_plain,
}

OCR_PARSERS = {
    ".jpg": _extract_ocr_tencent,
    ".jpeg": _extract_ocr_tencent,
    ".png": _extract_ocr_tencent,
    ".webp": _extract_ocr_tencent,
    ".gif": _extract_ocr_tencent,
}


def is_supported_document(path: Path) -> bool:
    return path.suffix.lower() in DOC_PARSERS


def is_supported_image_for_ocr(path: Path) -> bool:
    return path.suffix.lower() in OCR_PARSERS


def parse_asset(path: Path) -> dict:
    """解析/OCR 资产文件，返回 {text, text_length, text_status, ocr_used, error}。"""
    ext = path.suffix.lower()
    if not path.exists():
        return {"text": None, "text_status": "failed", "text_error": "file not found", "ocr_used": None}

    if path.stat().st_size > MAX_FILE_BYTES:
        return {
            "text": None,
            "text_status": "failed",
            "text_error": f"file too large ({path.stat().st_size} > {MAX_FILE_BYTES})",
            "ocr_used": None,
        }

    try:
        if ext in DOC_PARSERS:
            text, meta = DOC_PARSERS[ext](path)
            text = (text or "").strip()
            truncated = False
            if len(text) > MAX_TEXT_CHARS:
                text = text[:MAX_TEXT_CHARS]
                truncated = True
            return {
                "text": text,
                "text_length": len(text),
                "text_status": "ready",
                "text_error": ("truncated to 30k chars" if truncated else None),
                "ocr_used": None,
                "meta": {**meta, "truncated": truncated},
            }
        if ext in OCR_PARSERS:
            text, meta = OCR_PARSERS[ext](path)
            text = (text or "").strip()
            truncated = False
            if len(text) > MAX_TEXT_CHARS:
                text = text[:MAX_TEXT_CHARS]
                truncated = True
            return {
                "text": text,
                "text_length": len(text),
                "text_status": "ready",
                "text_error": ("truncated to 30k chars" if truncated else None),
                "ocr_used": meta.get("service", "tencent-ocr"),
                "meta": {**meta, "truncated": truncated},
            }
        return {"text": None, "text_status": "failed", "text_error": f"unsupported: {ext}", "ocr_used": None}
    except Exception as exc:  # noqa: BLE001
        logger.exception("Asset parse failed: %s", path)
        return {
            "text": None,
            "text_status": "failed",
            "text_error": f"{type(exc).__name__}: {str(exc)[:300]}",
            "ocr_used": None,
        }


def parse_bytes(data: bytes, filename: str) -> dict:
    """解析来自内存的字节（用于上传后立即解析）。"""
    ext = Path(filename or "").suffix.lower()
    if ext not in DOC_PARSERS and ext not in OCR_PARSERS:
        return {"text": None, "text_status": "failed", "text_error": f"unsupported: {ext}", "ocr_used": None}

    if len(data) > MAX_FILE_BYTES:
        return {
            "text": None,
            "text_status": "failed",
            "text_error": f"file too large ({len(data)} > {MAX_FILE_BYTES})",
            "ocr_used": None,
        }

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        return parse_asset(tmp_path)
    finally:
        try:
            tmp_path.unlink()
        except Exception:
            pass
